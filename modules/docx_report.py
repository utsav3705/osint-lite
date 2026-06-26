import os
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

def generate_docx(subject_info: dict, github_data: dict, username_data: list, whois_data: dict, risk_data: dict,
                  aliases: list, social_data: dict, news_data: dict, company_data: dict, ai_summary: str,
                  email_breaches: dict = None, output_dir: str = "generated_reports") -> str:
    """
    Generates a DOCX investigation report.
    Returns the filename of the generated document.
    """
    os.makedirs(output_dir, exist_ok=True)
    name = subject_info.get("name") or subject_info.get("username") or "Unknown"
    safe_name = "".join([c if c.isalnum() else "_" for c in name])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"OSINT_Pro_{safe_name}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    
    # Title
    doc.add_heading("OSINT-Pro Investigation Report", 0)
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    if ai_summary:
        doc.add_paragraph(ai_summary)
    else:
        doc.add_paragraph(f"Investigation completed for subject: {name}.")
        
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Overall Risk Score: {risk_data.get('score', 0)}/100")
    doc.add_paragraph(f"Risk Classification: {risk_data.get('classification', 'Unknown')}")
    
    # Subject Info
    doc.add_heading("Subject Information", level=1)
    for key, value in subject_info.items():
        if value:
            doc.add_paragraph(f"{key.capitalize()}: {value}")
            
    # Email Breaches
    if email_breaches and email_breaches.get("breached"):
        doc.add_heading("Email Breach Intelligence", level=1)
        doc.add_paragraph(f"Total Breaches Found: {email_breaches.get('count')}")
        p = doc.add_paragraph("Breaches: ")
        p.add_run(", ".join(email_breaches.get("breaches", [])))

    # GitHub Findings
    if github_data and github_data.get("username"):
        doc.add_heading("GitHub Findings", level=1)
        doc.add_paragraph(f"Username: {github_data.get('username')}")
        doc.add_paragraph(f"Public Repos: {github_data.get('public_repos', 0)}")
        doc.add_paragraph(f"Followers: {github_data.get('followers', 0)}")
        
    # Social Media & Media Findings
    if social_data:
        doc.add_heading("Social & Media Profiles", level=1)
        for platform, profile in social_data.items():
            if profile:
                doc.add_paragraph(f"{platform.capitalize()}: Profile identified")
                
    # Media / News Findings
    if news_data and news_data.get("articles"):
        doc.add_heading("Media Findings", level=1)
        for article in news_data.get("articles", [])[:5]:
            doc.add_paragraph(f"- {article.get('title', 'Unknown Title')} ({article.get('sentiment', 'Neutral')})")
            
    # Company Intelligence
    if company_data and company_data.get("name"):
        doc.add_heading("Company Intelligence", level=1)
        doc.add_paragraph(f"Company: {company_data.get('name')}")
        doc.add_paragraph(f"Industry: {company_data.get('industry', 'N/A')}")
        
    # Recommendations
    doc.add_heading("Recommendations", level=1)
    if risk_data.get("classification") == "High":
        doc.add_paragraph("- Escalate investigation to senior analyst.")
        doc.add_paragraph("- Request additional KYC documents.")
    else:
        doc.add_paragraph("- Proceed with standard due diligence monitoring.")
        
    doc.save(filepath)
    return filename
